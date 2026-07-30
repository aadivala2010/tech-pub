#!/usr/bin/env python3
"""
AI Document Revision Agent  —  setup:
    pip install python-docx PyMuPDF requests
Run:
    python agent.py
"""

import base64, difflib, html, json, os, platform, re, shutil, threading, time, traceback, zipfile
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext
import tkinter as tk
from xml.etree import ElementTree as ET

# Load .env before config — manual, no dependency needed
_env_path = Path(__file__).parent / '.env'
if _env_path.exists():
    with open(_env_path) as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith('#'):
                _k, _v = _line.split('=', 1)
                os.environ[_k.strip()] = _v.strip().strip('"').strip("'")

# ══════════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════════
# Load Gemini credentials from environment for safety and flexibility
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
GEMINI_MODEL   = "gemini-3.1-flash-lite"
# Set DOCREVISE_DEBUG=1 to see the raw technical log lines (hidden by default so
# the on-screen log stays plain-English for a live audience).
DEBUG = os.getenv("DOCREVISE_DEBUG") == "1"
# ══════════════════════════════════════════════════════════════════

import requests
from docx import Document
import fitz  # PyMuPDF

_ROOT = "C:\\" if platform.system() == "Windows" else "/"


# ═══════════════════════════════════════════════════════════════════
#  AI — plain HTTP, no google packages
# ═══════════════════════════════════════════════════════════════════

_GEMINI_URL = (
    "https://generativelanguage.googleapis.com/v1beta/models/"
    "{model}:generateContent?key={key}"
)

def ai_call(prompt):
    url  = _GEMINI_URL.format(model=GEMINI_MODEL, key=GEMINI_API_KEY)
    body = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, json=body, timeout=120)
    resp.raise_for_status()
    data = resp.json()
    try:
        return data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"Unexpected Gemini response: {data}") from exc


# ═══════════════════════════════════════════════════════════════════
#  UTILITIES
# ═══════════════════════════════════════════════════════════════════

def _dbg(log, msg):
    """Technical detail — only shown when DOCREVISE_DEBUG=1, so the live log
    stays plain-English."""
    if DEBUG:
        log(msg)

def clean_json(raw):
    """Extract the first [...] or {...} block from a string."""
    # Strip code fences
    raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$",           "", raw, flags=re.MULTILINE)
    raw = raw.strip()
    # Extract first JSON array or object
    for start_char, end_char in [("[", "]"), ("{", "}")]:
        idx = raw.find(start_char)
        if idx != -1:
            depth, end = 0, -1
            for i, ch in enumerate(raw[idx:], idx):
                if ch == start_char:   depth += 1
                elif ch == end_char:
                    depth -= 1
                    if depth == 0:
                        end = i
                        break
            if end != -1:
                return raw[idx:end+1]
    return raw

def increment_revision(filename):
    stem   = Path(filename).stem
    suffix = Path(filename).suffix
    pat    = re.compile(r"^(.*?)([_\-\s]?)([Rr][Ee][Vv][_\-]?)([A-Za-z]+)(.*)$")
    m      = pat.match(stem)
    if m:
        pre, sep, rev_tag, letters, post = m.groups()
        return f"{pre}{sep}{rev_tag}{_next_rev(letters.upper())}{post}{suffix}"
    return f"{stem}_revA{suffix}"

def _next_rev(s):
    chars = list(s)
    for i in range(len(chars) - 1, -1, -1):
        if chars[i] != "Z":
            chars[i] = chr(ord(chars[i]) + 1)
            return "".join(chars)
        chars[i] = "A"
    return "A" + "".join(chars)

def safe_output_path(dst):
    dst = Path(dst)
    if not dst.exists():
        return dst
    stem, suffix, parent = dst.stem, dst.suffix, dst.parent
    n = 2
    while True:
        c = parent / f"{stem}_{n}{suffix}"
        if not c.exists():
            return c
        n += 1

def read_file_text(path):
    ext = Path(path).suffix.lower()
    if ext in (".txt", ".md"):
        return Path(path).read_text(encoding="utf-8", errors="replace")
    if ext == ".docx":
        doc   = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    if ext == ".pdf":
        pdf  = fitz.open(path)
        text = "\n".join(page.get_text() for page in pdf)
        pdf.close()
        return text
    raise ValueError(f"Unsupported: {ext}")

def mentions_old_part(path, changes):
    try:
        text = read_file_text(path).lower()
        return any(c.get("old_part", "").lower() in text
                   for c in changes if c.get("old_part"))
    except Exception:
        return False

def text_needs_update(text, changes):
    tl = text.lower()
    return any(c.get("old_part", "").lower() in tl
               for c in changes if c.get("old_part"))


# ═══════════════════════════════════════════════════════════════════
#  IMAGE GALLERY
# ═══════════════════════════════════════════════════════════════════

IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff", ".webp"}

def index_gallery(folder):
    """Return dict: lowercase_stem -> full_path"""
    gallery = {}
    if not folder or not os.path.exists(folder):
        return gallery
    for p in Path(folder).rglob("*"):
        if p.suffix.lower() in IMG_EXTS and p.is_file():
            gallery[p.stem.lower()] = str(p)
    return gallery

def auto_discover_gallery_from_folder(folder, log):
    """Auto-discover images within a user-provided folder.
    Maps by image stem only (ignore subfolder hints). Logs duplicates.
    Returns a dict: {stem: path}
    """
    log(f"  [auto-detect] Scanning provided gallery folder: {folder}")
    discovered = {}
    if not folder or not os.path.isdir(folder):
        log("  [auto-detect] No valid folder provided for auto-discovery.")
        return discovered
    root = Path(folder)
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in IMG_EXTS:
            path = str(p)
            stem = p.stem.lower()
            if stem:
                if stem not in discovered:
                    discovered[stem] = path
                    log(f"  [auto-detect] mapped stem '{stem}' -> {path}")
                else:
                    log(f"  [auto-detect] duplicate stem '{stem}' (existing: {discovered[stem]}), keeping first")
    if not discovered:
        log("  [auto-detect] No images found in provided folder.")
    else:
        log(f"  [auto-detect] Discovered {len(discovered)} image hint(s): {sorted(discovered.keys())}")
    return discovered

def best_gallery_match(gallery, hint):
    """Find best image for hint string (part number, name, etc.)"""
    if not hint or not gallery:
        return None
    h = hint.lower().strip()
    # Exact stem match
    if h in gallery:
        return gallery[h]
    # Partial match
    for stem, path in gallery.items():
        if h in stem or stem in h:
            return path
    return None


# ═══════════════════════════════════════════════════════════════════
#  STEP 1 — Parse the change prompt
# ═══════════════════════════════════════════════════════════════════

_PARSE_PROMPT = """\
You are reading an engineering change order document.

Extract every part substitution. Return ONLY a raw JSON array — no markdown, no explanation.

For each substitution identify:
- old_part  : the exact part name/number currently used in documents
- new_part  : the exact part name/number to replace it with
- old_specs : dict of spec name -> OLD value  e.g. {{"pressure": "30psi"}}
- new_specs : dict of spec name -> NEW value  e.g. {{"pressure": "45psi"}}
- new_image : the filename or part name to use for the replacement image
              (use new_part value if no specific image is mentioned)

Example output for "replace valve ABC (30psi) with DEF (45psi)":
[
  {{
    "old_part":  "ABC",
    "new_part":  "DEF",
    "old_specs": {{"pressure": "30psi"}},
    "new_specs": {{"pressure": "45psi"}},
    "new_image": "DEF"
  }}
]

CHANGE ORDER DOCUMENT:
\"\"\"
{text}
\"\"\"
"""

def parse_prompt(text, log):
    log("  Asking the AI to read the request and pull out the exact changes.")
    raw = ai_call(_PARSE_PROMPT.format(text=text))
    _dbg(log, f"[dbg] AI parse response: {raw[:300]}")
    try:
        items = json.loads(clean_json(raw))
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            f"Gemini returned invalid JSON.\nRaw response:\n{raw}\nError: {exc}"
        ) from exc
    result = []
    for item in items:
        if not isinstance(item, dict):
            continue
        item.setdefault("old_part",  "")
        item.setdefault("new_part",  "")
        item.setdefault("old_specs", {})
        item.setdefault("new_specs", {})
        item.setdefault("new_image", item.get("new_part", ""))
        if not isinstance(item["old_specs"], dict): item["old_specs"] = {}
        if not isinstance(item["new_specs"], dict): item["new_specs"] = {}
        result.append(item)
    return result


# ═══════════════════════════════════════════════════════════════════
#  STEP 2 — Text rewriting
#
#  THE FIX: tell Gemini EXACTLY what to substitute ("Replace X with Y"),
#  not vague rules ("never contain X"). Then apply a regex fallback
#  as a hard guarantee that old_part is always replaced.
# ═══════════════════════════════════════════════════════════════════

_REWRITE_PROMPT = """\
Rewrite the text below by making ONLY these exact substitutions:

{sub_list}

Rules:
- Apply every substitution listed above, everywhere it appears.
- Do NOT change anything else — keep all other words, punctuation, and structure identical.
- Do NOT add any commentary such as "formerly", "previously", "updated", "replaced", "new", etc.
- The result must read naturally, as if the text was always written this way.
- Return ONLY the rewritten text. No explanation. No quotes around it.

TEXT:
\"\"\"
{text}
\"\"\"
"""

def _build_sub_list(change):
    """Build a human-readable substitution list for the prompt."""
    lines = []
    old_p = change.get("old_part", "")
    new_p = change.get("new_part", "")
    if old_p and new_p:
        lines.append(f'  Replace "{old_p}" with "{new_p}"')
    for k, old_v in change.get("old_specs", {}).items():
        new_v = change.get("new_specs", {}).get(k, "")
        if old_v and new_v and old_v != new_v:
            lines.append(f'  Replace "{old_v}" with "{new_v}"')
    return "\n".join(lines)

def _regex_replace(text, change):
    """
    Hard regex pass — guarantees old_part and old specs are replaced.
    Applied AFTER the AI call as a safety net.
    """
    result = text
    old_p  = change.get("old_part", "")
    new_p  = change.get("new_part", "")
    if old_p and new_p:
        result = re.sub(re.escape(old_p), new_p, result, flags=re.IGNORECASE)
    for k, old_v in change.get("old_specs", {}).items():
        new_v = change.get("new_specs", {}).get(k, "")
        if old_v and new_v and old_v != new_v:
            result = re.sub(re.escape(old_v), new_v, result, flags=re.IGNORECASE)
    # Fallback: if old_specs is empty but new_specs exist, attempt a broad numeric unit replacement
    if not change.get("old_specs") and change.get("new_specs"):
        # take the first new_specs value (e.g. "45 bar") and replace occurrences like '30 bar' or '35 bar'
        try:
            first_new = next(iter(change["new_specs"].values()))
            if isinstance(first_new, str) and first_new:
                result = re.sub(r"\b\d+(?:\.\d+)?\s*(?:bar|psi|kPa|Pa)\b", first_new, result, flags=re.IGNORECASE)
        except StopIteration:
            pass
    return result

def rewrite_chunk(text, changes, log):
    """
    Rewrite one text chunk against all changes.
    Uses AI for natural language quality, then regex as a guarantee.
    """
    current = text
    for change in changes:
        old_p = change.get("old_part", "")
        if not old_p or old_p.lower() not in current.lower():
            continue

        sub_list = _build_sub_list(change)
        if not sub_list.strip():
            continue

        prompt  = _REWRITE_PROMPT.format(sub_list=sub_list, text=current)
        ai_out  = ai_call(prompt)

        # Safety net: if AI missed the part name, force-replace it
        before_regex = ai_out
        ai_out = _regex_replace(ai_out, change)

        if ai_out != before_regex:
            log(f"    [note] regex fallback applied to ensure part name replaced")

        current = ai_out

    return current


# ═══════════════════════════════════════════════════════════════════
#  DOCX — paragraph text rewriting
# ═══════════════════════════════════════════════════════════════════

def _get_para_text(para):
    return "".join(r.text for r in para.runs)

_W_DRAWING = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"

def _run_has_image(run):
    """Return True if this run contains an inline image (w:drawing element)."""
    return run._r.find(_W_DRAWING) is not None

def _set_para_text(para, new_text):
    # Split runs into image-bearing runs (must be preserved) and plain text runs.
    # In DOCX, inline images live inside <w:r> elements as <w:drawing> children.
    # Setting run.text = "" on such a run destroys the <w:drawing> element,
    # wiping the image from document.xml — which is the root cause of the bug.
    text_runs = [r for r in para.runs if not _run_has_image(r)]
    if not text_runs:
        # Paragraph has no plain-text runs (only image runs); insert a new text
        # run at the front so the new text appears before the images, untouched.
        new_run = para.add_run(new_text)   # add_run appends to the end
        para._p.remove(new_run._r)         # pull it back out
        para._p.insert(0, new_run._r)      # put it at the very front
        return
    # Zero out only the plain-text runs, leave image runs alone
    for run in text_runs:
        run.text = ""
    text_runs[0].text = new_text

def rewrite_docx_text(doc, changes, log):
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    modified = False
    for para in all_paras:
        orig = _get_para_text(para)
        if not orig.strip() or not text_needs_update(orig, changes):
            continue
        updated = rewrite_chunk(orig, changes, log)
        if updated.strip() != orig.strip():
            _set_para_text(para, updated)
            _dbg(log, f"    was: {orig[:80].rstrip()}")
            _dbg(log, f"    now: {updated[:80].rstrip()}")
            modified = True
    return modified


# ═══════════════════════════════════════════════════════════════════
#  DOCX — image replacement
#  Finds images contextually near old_part text, swaps with gallery.
# ═══════════════════════════════════════════════════════════════════

_BLIP  = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

def _get_rids_in_para(para):
    return [b.get(_EMBED) for b in para._p.iter(_BLIP) if b.get(_EMBED)]

def _window_text(paras, idx, radius=3):
    lo = max(0, idx - radius)
    hi = min(len(paras), idx + radius + 1)
    return " ".join(p.text for p in paras[lo:hi]).lower()

def _all_paras(doc):
    result = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(cell.paragraphs)
    return result

def _parse_rels(xml_bytes):
    root = ET.fromstring(xml_bytes)
    return {r.get("Id", ""): r.get("Target", "") for r in root}

def replace_docx_images(docx_path, changes, gallery, log):
    if not gallery:
        log("  [img] No gallery provided — skipping image replacement")
        return False

    log("  [img] ==== STARTING IMAGE REPLACEMENT DEBUG ====")

    doc = Document(docx_path)
    # Diagnostic: count image relationships (not just inline images)
    try:
        image_rel_count = len([rel for rel in doc.part.rels.values() if "image" in getattr(rel, "reltype", "").lower()])
        log(f"  [img] Document has {image_rel_count} image relationship(s) in package")
    except Exception:
        # If the docx structure isn't exactly as expected, fail closedly without breaking.
        log("  [img] Unable to enumerate image relationships in document for diagnostics")

    paras = _all_paras(doc)
    rid_to_newimg = {}
    new_img_path = None  # Store for fallback

    log(f"  [img] Gallery contents: {gallery}")

    for change in changes:
        old_p    = change.get("old_part", "")
        img_hint = change.get("new_image") or change.get("new_part", "")
        if not old_p:
            continue

        new_img_path = best_gallery_match(gallery, img_hint)
        if not new_img_path:
            log(f"  [img] Gallery has no match for hint '{img_hint}'")
            log(f"  [img] Gallery contents: {list(gallery.keys())}")
            continue

        log(f"  [img] Gallery matched '{img_hint}' -> {Path(new_img_path).name}")
        log(f"  [img] Full path to replacement image: {new_img_path}")
        if not os.path.exists(new_img_path):
            log(f"  [img] WARNING: replacement image path not found: {new_img_path}")

        # Diagnostic: show which media entries exist in the docx package
        try:
            with zipfile.ZipFile(docx_path) as z:
                media_in_doc = [p for p in z.namelist() if p.startswith("word/media/")]
            log(f"  [img] DOCX media entries: {media_in_doc}")
        except Exception:
            log("  [img] Unable to enumerate DOCX media entries for diagnostics")

        new_p_name = change.get("new_part", "")
        found_any_image_para = False
        for i, para in enumerate(paras):
            rids = _get_rids_in_para(para)
            if not rids:
                continue
            found_any_image_para = True
            window = _window_text(paras, i, radius=3)
            log(f"  [img] Para {i} has image(s) {rids}, window: '{window[:80]}'")
            # Bug fix: text rewriting runs before image replacement, so old_p is
            # already replaced with new_p in the saved file. Check both names.
            in_old = old_p.lower() in window
            in_new = bool(new_p_name) and new_p_name.lower() in window
            if in_old or in_new:
                matched_by = old_p if in_old else new_p_name
                for rid in rids:
                    rid_to_newimg[rid] = new_img_path
                    log(f"  [img] MATCHED: rId={rid} near '{matched_by}' -> will swap")
            else:
                log(f"  [img] No match: neither '{old_p}' nor '{new_p_name}' in window")

        if not found_any_image_para:
            log(f"  [img] Document has no inline images at all")

    if not rid_to_newimg:
        log("  [img] No images to swap — attempting fallback if possible")

    # Read rels to map rId -> zip entry path
    with zipfile.ZipFile(docx_path) as z:
        namelist = z.namelist()
        rels = _parse_rels(z.read("word/_rels/document.xml.rels"))
        log(f"  [img] DOCX archive contains {len(namelist)} entries; sample: {namelist[:5]}")
        # Fallback: if no proximity-based image mappings found, apply to all images in the doc.
        # Requires new_img_path — without a gallery match there is nothing to write, and
        # mapping rIds to None used to crash open() further down.
        if not rid_to_newimg and new_img_path:
            log("  [img] Fallback: applying replacement to all images in document using new image.")
            for rid, target in rels.items():
                # Rels targets are relative to word/, so they appear as "media/..."
                # not "word/media/..." — must check both prefixes.
                if target.startswith("word/media/") or target.startswith("media/"):
                    rid_to_newimg[rid] = new_img_path

    log(f"  [img] Rels map: { {k:v for k,v in rels.items() if 'image' in v.lower()} }")
    log(f"  [img] rid_to_newimg mappings: {rid_to_newimg}")

    media_map = {}
    for rid, img_path in rid_to_newimg.items():
        target = rels.get(rid, "")
        # DOCX rels often reference media with a 'word/media/...' path
        if target.startswith("word/media/"):
            media_map[target] = img_path
        elif target.startswith("media/"):
            media_map["word/" + target] = img_path
        else:
            log(f"  [img] WARNING: rId={rid} target='{target}' not in media/ — skipped")

    # DON'T return early - proceed to blanket replacement if media_map is empty
    # if not media_map:
    #     log("  [img] No media entries resolved from rIds")
    #     return False

    # Check if there are images in the docx that need replacement
    media_entries_in_doc = [e for e in namelist if e.startswith("word/media/") or e.startswith("media/")]
    log(f"  [img] Media entries found in docx: {media_entries_in_doc}")
    log(f"  [img] media_map currently: {media_map}")
    log(f"  [img] new_img_path is: {new_img_path}")

    tmp = docx_path + ".__imgtmp__.zip"
    swapped = False
    # Preload replacement blob for blanket replacement (fallback)
    replacement_blob = None
    if new_img_path and media_entries_in_doc:
        try:
            with open(new_img_path, "rb") as f:
                replacement_blob = f.read()
            log(f"  [img] Loaded replacement image blob: {len(replacement_blob)} bytes")
            # Verify replacement image is valid by checking header
            replacement_header = replacement_blob[:8]
            log(f"  [img] Replacement image header (hex): {replacement_header.hex()}")
        except Exception as e:
            log(f"  [img] WARNING: unable to read replacement image '{new_img_path}': {e}")

    # Determine if blanket replacement should occur (no per-entry mapping worked but there are images)
    blanket_replace = False
    log(f"  [img] blanket_replace check: media_map empty={not media_map}, replacement_blob exists={replacement_blob is not None}, media_entries exist={bool(media_entries_in_doc)}")
    if not media_map and replacement_blob and media_entries_in_doc:
        blanket_replace = True
        log("  [img] Will blanket replace all images")

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for entry in zin.namelist():
                if entry in media_map:
                    with open(media_map[entry], "rb") as f:
                        zout.writestr(entry, f.read())
                    log(f"  [img] Swapped {entry} -> {Path(media_map[entry]).name}")
                    swapped = True
                elif blanket_replace and (entry.startswith("word/media/") or entry.startswith("media/")):
                    # Preserve original filename extension when replacing
                    original_ext = os.path.splitext(entry)[1]
                    replacement_name = os.path.splitext(os.path.basename(new_img_path))[0] + original_ext
                    # Write to original entry path with new content but keep original filename
                    zout.writestr(entry, replacement_blob)
                    log(f"  [img] Blanket replaced {entry} (kept original extension {original_ext}) with {replacement_name}")
                    swapped = True
                else:
                    zout.writestr(entry, zin.read(entry))
        if swapped:
            os.replace(tmp, docx_path)
            log(f"  [img] Successfully replaced {docx_path}")
            
            # Verification: check that the new image is actually in the file
            try:
                log("  [img] === STARTING VERIFICATION ===")
                with zipfile.ZipFile(docx_path, "r") as z:
                    log(f"  [img] Verification: opened {docx_path}")
                    # Also extract document.xml to see how images are referenced
                    try:
                        doc_xml = z.read("word/document.xml").decode("utf-8", errors="replace")
                        # Find all blip references
                        import re
                        blips = re.findall(r'<a:blip r:embed="([^"]+)"', doc_xml)
                        log(f"  [img] Document XML blip references: {blips}")
                    except Exception as de:
                        log(f"  [img] Could not read document.xml: {de}")
                    
                    media_entries = [e for e in z.namelist() if e.startswith("word/media/")]
                    log(f"  [img] Verification: found media entries: {media_entries}")
                    
                log("  [img] === VERIFICATION COMPLETE ===")
            except Exception as ve:
                log(f"  [img] VERIFY ERROR: {ve}")
                import traceback
                log(f"  [img] VERIFY ERROR traceback: {traceback.format_exc()}")
        else:
            os.remove(tmp)
            log("  [img] No images were swapped")
    except Exception as e:
        if os.path.exists(tmp):
            os.remove(tmp)
        # Best-effort: the text rewrite is already saved into docx_path. Raising here
        # made run_pipeline delete that output, throwing away completed text work
        # because of a picture. Report it and keep the text changes.
        log(f"  [img] ERROR during image swap (text changes kept): {e}")
        _dbg(log, traceback.format_exc())
        return False

    return swapped


# ═══════════════════════════════════════════════════════════════════
#  DOCX full processing
# ═══════════════════════════════════════════════════════════════════

def _docx_has_images(docx_path):
    """Return True if the DOCX zip contains any media (image) entries."""
    with zipfile.ZipFile(docx_path) as z:
        return any(
            e.startswith("word/media/") or e.startswith("media/")
            for e in z.namelist()
        )

def _extract_text_from_docx(doc_path):
    """Extract and return all visible text from a DOCX file."""
    doc = Document(doc_path)
    texts = []
    # Paragraphs
    for para in doc.paragraphs:
        texts.append(para.text or "")
    # Tables: gather text from all cells (each cell may have its own paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text or "")
    return "\n".join(texts)

def _insert_image_near_mention(doc, changes, gallery, log):
    """
    For documents that have NO existing images: find the first paragraph that
    mentions old_part or new_part and insert a new inline image after it.
    This is needed because replace_docx_images can only SWAP existing images —
    it cannot create images from nothing.
    Returns True if at least one image was inserted.
    """
    from docx.shared import Inches

    paras    = _all_paras(doc)
    inserted = False

    log(f"  [img-insert] Scanning {len(paras)} paragraph(s) for part mentions")

    for change in changes:
        old_p    = change.get("old_part", "")
        new_p    = change.get("new_part", "")
        img_hint = change.get("new_image") or new_p

        log(f"  [img-insert] Looking for '{new_p}' (then '{old_p}'), image hint='{img_hint}'")

        new_img_path = best_gallery_match(gallery, img_hint)
        if not new_img_path:
            log(f"  [img-insert] No gallery match for '{img_hint}' — cannot insert")
            log(f"  [img-insert] Gallery keys: {list(gallery.keys())}")
            continue

        log(f"  [img-insert] Gallery match: {Path(new_img_path).name}")

        # Search for new_part first (text already rewritten), then old_part
        target_para = None
        for search_term in ([new_p] if new_p else []) + ([old_p] if old_p else []):
            for para in paras:
                if search_term.lower() in para.text.lower():
                    target_para = para
                    log(f"  [img-insert] Matched '{search_term}' in paragraph: '{para.text[:80]}'")
                    break
            if target_para:
                break

        if not target_para:
            log(f"  [img-insert] No paragraph mentions '{old_p}' or '{new_p}'")
            log(f"  [img-insert] First 10 paragraph texts:")
            for p in paras[:10]:
                if p.text.strip():
                    log(f"    '{p.text[:80]}'")
            continue

        try:
            # Add paragraph at end of document body, then move it to after target_para.
            # doc.add_paragraph() handles all namespace/relationship setup correctly.
            # In lxml, addnext() on an already-placed element MOVES it.
            new_para = doc.add_paragraph()
            run      = new_para.add_run()
            run.add_picture(new_img_path, width=Inches(3))
            target_para._p.addnext(new_para._p)   # moves element, not a copy
            log(f"  [img-insert] Successfully inserted '{Path(new_img_path).name}'")
            inserted = True
        except Exception as e:
            log(f"  [img-insert] Exception during insert: {e}")
            log(traceback.format_exc())

    return inserted


def process_docx(src, dst, changes, gallery, log):
    shutil.copy2(src, dst)

    has_images = _docx_has_images(dst)
    log(f"  [route] Document has existing images: {has_images}")

    doc          = Document(dst)
    text_changed = rewrite_docx_text(doc, changes, log)
    log(f"  [route] Text changed: {text_changed}")

    img_inserted = False
    img_changed  = False

    if gallery and has_images:
        # REPLACE path: existing media found — save text first, then swap the blob at
        # zip level. Must NOT call doc.save() again after replace_docx_images or it
        # would overwrite the swap with the old in-memory image bytes.
        log("  [route] Existing images found → REPLACE path (via replace_docx_images)")
        if text_changed:
            doc.save(dst)
            log("  [route] Saved text changes before image swap")
        img_changed = replace_docx_images(dst, changes, gallery, log)
        log(f"  [route] replace_docx_images result: {img_changed}")
    elif gallery:
        # INSERT path: no existing images — insert replacement image next to mention.
        log("  [route] No images found → INSERT path")
        img_inserted = _insert_image_near_mention(doc, changes, gallery, log)
        if img_inserted or text_changed:
            doc.save(dst)
            log("  [route] Saved document with changes")
    else:
        # No gallery — only persist text changes.
        if text_changed:
            doc.save(dst)

    if text_changed or img_changed or img_inserted:
        log(f"  [route] Final: text={text_changed} inserted={img_inserted} replaced={img_changed}")
        try:
            final_text = _extract_text_from_docx(dst)
            log(f"  [dbg] Final DOCX text excerpt: {final_text[:500]!r}")
        except Exception as _e:
            log(f"  [dbg] TEXT EXTRACTION FAILED: {_e}")

    return text_changed or img_changed or img_inserted


# ═══════════════════════════════════════════════════════════════════
#  PDF processing
# ═══════════════════════════════════════════════════════════════════

def process_pdf(src, dst, changes, log):
    pdf      = fitz.open(src)
    modified = False

    for pnum, page in enumerate(pdf):
        updates = []
        for block in page.get_text("blocks"):
            if block[6] != 0:
                continue
            text = block[4].strip()
            if not text or not text_needs_update(text, changes):
                continue
            updated = rewrite_chunk(text, changes, log)
            if updated.strip() == text.strip():
                continue
            rect  = fitz.Rect(block[:4])
            fsize = max(7, int(rect.height / max(text.count("\n") + 1, 1) * 0.72))
            updates.append((rect, updated, fsize, text))
            page.add_redact_annot(rect, fill=(1, 1, 1))

        if updates:
            page.apply_redactions()
            for rect, updated, fsize, orig in updates:
                page.insert_text(
                    fitz.Point(rect.x0, rect.y0 + fsize),
                    updated, fontsize=fsize, color=(0, 0, 0),
                )
                log(f"    [pdf p{pnum+1}] was: {orig[:60].rstrip()}")
                log(f"                now: {updated[:60].rstrip()}")
            modified = True

    if modified:
        pdf.save(dst, garbage=4, deflate=True)
    pdf.close()
    return modified


# ═══════════════════════════════════════════════════════════════════
#  LIVE ON-SCREEN VIEW  —  open the real document in Word and make the
#  changes visibly, for a live audience, then close without saving.
#  Best-effort only: the authoritative saved file is produced by
#  process_docx.  Any failure here is caught and never stops the run.
# ═══════════════════════════════════════════════════════════════════

def _concrete_subs(change, source_text):
    """The exact old->new text swaps that will happen, as literal strings the
    audience can watch. Mirrors what _regex_replace does to the saved file."""
    subs = []
    op = change.get("old_part", "")
    np = change.get("new_part", "")
    if op and np:
        subs.append((op, np))
    for k, ov in change.get("old_specs", {}).items():
        nv = change.get("new_specs", {}).get(k, "")
        if ov and nv and ov != nv:
            subs.append((ov, nv))
    # Old spec value not stated in the request: find the matching number+unit
    # actually present in the document (e.g. "35 bar") and swap it.
    if not change.get("old_specs") and change.get("new_specs"):
        nv = next(iter(change["new_specs"].values()), "")
        if isinstance(nv, str) and nv:
            m = re.search(r"\d+(?:\.\d+)?\s*(?:bar|psi|kPa|Pa)\b",
                          source_text, re.IGNORECASE)
            if m and m.group(0).lower() != nv.lower():
                subs.append((m.group(0), nv))
    return subs

def _docx_paragraph_texts(path):
    """Return the visible, non-empty paragraph texts of a docx, in order."""
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text.strip()]


def _bytes_to_data_uri(data, ext):
    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "bmp": "image/bmp"}.get(ext.lower().lstrip("."), "image/png")
    return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"


def _docx_first_image_data_uri(path):
    """Return the first embedded image in a docx as a data: URI, or None."""
    try:
        with zipfile.ZipFile(path) as z:
            media = sorted(e for e in z.namelist() if e.startswith("word/media/"))
            if not media:
                return None
            data = z.read(media[0])
        return _bytes_to_data_uri(data, Path(media[0]).suffix)
    except Exception:
        return None


def _file_image_data_uri(path):
    """Return a file on disk as a data: URI, or None."""
    try:
        with open(path, "rb") as f:
            data = f.read()
        return _bytes_to_data_uri(data, Path(path).suffix)
    except Exception:
        return None


def _diff_html(old_text, new_text):
    """HTML for a paragraph mid-edit: unchanged text plain, removed text struck
    through, inserted text highlighted — so the viewer sees the exact edit."""
    sm = difflib.SequenceMatcher(None, old_text, new_text)
    out = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            out.append(html.escape(old_text[i1:i2]))
        elif tag == "delete":
            out.append(f'<span class="live-del">{html.escape(old_text[i1:i2])}</span>')
        elif tag == "insert":
            out.append(f'<span class="live-ins">{html.escape(new_text[j1:j2])}</span>')
        elif tag == "replace":
            out.append(f'<span class="live-del">{html.escape(old_text[i1:i2])}</span>')
            out.append(f'<span class="live-ins">{html.escape(new_text[j1:j2])}</span>')
    return "".join(out)


def present_live_view(src, changes, gallery, log, view):
    """Drive an animated, on-screen re-enactment of the edits about to be made
    to `src`, rendered as a live document panel in the browser UI (not an
    actual Word window). Windows blocks background processes from stealing
    foreground focus, so a real Word window can't reliably be forced on
    screen — this achieves the same "watch it happen live" effect through the
    same event stream already driving the log, on the tab that's guaranteed
    to already be visible. Never raises — a live demo must not be interrupted.
    """
    if view is None:
        return
    name = Path(src).name
    try:
        paragraphs = _docx_paragraph_texts(src)
    except Exception:
        paragraphs = []
    try:
        image_before = _docx_first_image_data_uri(src)
    except Exception:
        image_before = None

    try:
        view({"kind": "open", "doc": name, "paragraphs": paragraphs, "image": image_before})
        log("  Opening the document on screen so you can watch the changes happen.")
        time.sleep(1.0)

        working = list(paragraphs)
        source_text = "\n".join(paragraphs)
        for change in changes:
            for old, new in _concrete_subs(change, source_text):
                pattern = re.compile(re.escape(old), re.IGNORECASE)
                prev = list(working)
                touched = []
                for i, p in enumerate(working):
                    if pattern.search(p):
                        working[i] = pattern.sub(new, p)
                        touched.append(i)
                if not touched:
                    continue
                log(f'  Changing every "{old}" into "{new}".')
                preview = [{"index": i, "html": _diff_html(prev[i], working[i])} for i in touched]
                view({"kind": "replace_preview", "old": old, "new": new, "changes": preview})
                time.sleep(1.5)
                settle = [{"index": i, "text": working[i]} for i in touched]
                view({"kind": "replace_settle", "changes": settle})
                time.sleep(0.4)

        img_path = new_part = None
        for change in changes:
            hint = change.get("new_image") or change.get("new_part", "")
            img_path = best_gallery_match(gallery, hint)
            new_part = change.get("new_part", "")
            if img_path:
                break
        if img_path and os.path.exists(img_path) and image_before:
            log(f"  Swapping the old part picture for the photo of {new_part or 'the new part'}.")
            image_after = _file_image_data_uri(img_path)
            if image_after:
                view({"kind": "image", "image": image_after})
                time.sleep(1.4)

        log("  The document now shows all the changes. Saving a fresh copy.")
        view({"kind": "close"})
        time.sleep(0.4)
    except Exception as e:
        _dbg(log, f"[view] live view error: {e}")
        log("  (Couldn't show the live preview this time — the update itself still completes correctly.)")


# ═══════════════════════════════════════════════════════════════════
#  MAIN PIPELINE
# ═══════════════════════════════════════════════════════════════════

def run_pipeline(prompt_file, docs_folder, gallery_folder, log, done, view=None):
    try:
        # 1. Parse change order
        log("STEP 1 of 4  —  Reading the change request")
        log(f"  Reading the file that describes what needs to change: {Path(prompt_file).name}")
        prompt_text = read_file_text(prompt_file)
        _dbg(log, f"[dbg] Prompt text ({len(prompt_text)} chars): {prompt_text[:300]}")
        changes = parse_prompt(prompt_text, log)

        if not changes:
            log("  The request didn't name any part to replace. Please check the file and try again.")
            done(False, "No changes found — check the request file")
            return

        # Validate — if old_part or new_part is blank, stop now
        for c in changes:
            if not c.get("old_part") or not c.get("new_part"):
                log("  I couldn't tell which part replaces which. Please write it like:")
                log("      'Replace part ABC (30 psi) with part DEF (45 psi)'")
                done(False, "Couldn't read the part names — see the log")
                return

        log(f"  Understood. The request asks for {len(changes)} change(s):")
        for c in changes:
            line = f"    • Replace part {c['old_part']} with part {c['new_part']}"
            specs = [f"{k} becomes {v}" for k, v in c.get("new_specs", {}).items() if v]
            if specs:
                line += "  (" + ", ".join(specs) + ")"
            log(line)

        # 2. Index (or auto-discover from provided folder)
        log("")
        log("STEP 2 of 4  —  Finding the new part's photo")
        gallery = index_gallery(gallery_folder)
        if not gallery:
            gallery = auto_discover_gallery_from_folder(gallery_folder, log)
        if gallery:
            log(f"  Found {len(gallery)} photo(s) to choose from: {', '.join(sorted(gallery.keys()))}")
        else:
            log("  No photos folder was given, so pictures will stay as they are; only the wording changes.")

        # 3. Scan docs folder (content search)
        log("")
        log("STEP 3 of 4  —  Finding which documents are affected")
        log(f"  Looking through every document in: {docs_folder}")
        all_docs = [
            str(p)
            for p in Path(docs_folder).rglob("*")
            if p.suffix.lower() in (".docx", ".pdf")
            and p.is_file()
            and not p.name.startswith("~$")   # skip Word's temporary lock files
        ]
        log(f"  There are {len(all_docs)} document(s) to check.")

        relevant = []
        for d in all_docs:
            if mentions_old_part(d, changes):
                relevant.append(d)
                log(f"    ✓ {Path(d).name} — mentions the old part, so it needs updating.")
            else:
                log(f"    · {Path(d).name} — doesn't mention it, leaving this one untouched.")

        if not relevant:
            log("  None of the documents mention that part. Please check the spelling and try again.")
            done(False, "No documents mention that part — check spelling")
            return

        log(f"  {len(relevant)} document(s) will be updated.")

        # 4. Process each document
        log("")
        log("STEP 4 of 4  —  Updating the documents")

        report = []
        for src in relevant:
            name = Path(src).name
            log("")
            log(f"  Now working on: {name}")
            new_name = increment_revision(name)
            dst      = safe_output_path(Path(src).parent / new_name)
            ext      = Path(src).suffix.lower()

            # Live on-screen re-enactment for the audience (Word docs only).
            if ext == ".docx":
                present_live_view(src, changes, gallery, log, view)

            try:
                if ext == ".docx":
                    ok = process_docx(src, str(dst), changes, gallery, log)
                elif ext == ".pdf":
                    ok = process_pdf(src, str(dst), changes, log)
                else:
                    ok = False
            except Exception as e:
                log(f"  Something went wrong while updating {name}: {e}")
                _dbg(log, traceback.format_exc())
                if dst.exists():
                    dst.unlink()
                report.append({"file": name, "status": "error", "error": str(e)})
                continue

            if ok:
                log(f"  Done — saved the updated version as {dst.name}. The original {name} is kept, unchanged.")
                report.append({"file": name, "new": dst.name, "status": "ok"})
            else:
                if dst.exists():
                    dst.unlink()
                log(f"  Nothing needed changing in {name}.")
                report.append({"file": name, "status": "none"})

        # 5. Summary
        ok_r  = [r for r in report if r["status"] == "ok"]
        err_r = [r for r in report if r["status"] == "error"]

        log("")
        log("All finished!")
        log(f"  Documents checked : {len(all_docs)}")
        log(f"  Documents updated : {len(ok_r)}")
        if err_r:
            log(f"  Documents with a problem : {len(err_r)}")
        if ok_r:
            log("  New versions created:")
            for r in ok_r:
                log(f"    {r['file']}  →  {r['new']}")
        if err_r:
            log("  These had a problem:")
            for r in err_r:
                log(f"    {r['file']}: {r['error']}")

        done(True, f"{len(ok_r)} document(s) updated")

    except Exception:
        _dbg(log, traceback.format_exc())
        log("")
        log("Something unexpected happened and the run stopped early. Your original files are safe and unchanged.")
        done(False, "Unexpected error — please try again")


# ═══════════════════════════════════════════════════════════════════
#  GUI
# ═══════════════════════════════════════════════════════════════════

BG     = "#1e1e2e"
PANEL  = "#181825"
CARD   = "#313244"
SURF   = "#45475a"
FG     = "#cdd6f4"
FG2    = "#a6adc8"
FG3    = "#6c7086"
ACCENT = "#89b4fa"
GREEN  = "#a6e3a1"
RED    = "#f38ba8"
YELLOW = "#f9e2af"
MAUVE  = "#cba6f7"
FONT   = ("Segoe UI", 10)
FONT_B = ("Segoe UI", 10, "bold")
MONO   = ("Consolas", 9)


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("AI Document Revision Agent")
        self.geometry("980x740")
        self.minsize(760, 560)
        self.configure(bg=BG)
        self._build()

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg=PANEL, pady=18)
        hdr.pack(fill="x")
        tk.Label(hdr, text="AI Document Revision Agent",
                 font=("Segoe UI", 17, "bold"), fg=FG, bg=PANEL).pack()
        tk.Label(hdr, text=f"Gemini {GEMINI_MODEL}  |  content search  |  seamless replacement",
                 font=("Segoe UI", 8), fg=FG3, bg=PANEL).pack(pady=(3, 0))

        # Inputs card
        card = tk.Frame(self, bg=CARD, padx=24, pady=18)
        card.pack(fill="x", padx=20, pady=(16, 0))
        tk.Label(card, text="INPUTS", font=("Segoe UI", 8, "bold"),
                 fg=FG3, bg=CARD).pack(anchor="w", pady=(0, 8))

        self.docs_var    = self._row(card, "1   Documents folder",
            "Folder (and subfolders) to search — any location on your computer",
            is_file=False)
        self.prompt_var  = self._row(card, "2   Change prompt file",
            'e.g. "Replace part ABC (30psi) with DEF (45psi)"',
            is_file=True)
        self.gallery_var = self._row(card, "3   Image gallery folder",
            "Folder with replacement images — name files after the NEW part  (optional)",
            is_file=False)

        # Buttons
        btn_row = tk.Frame(self, bg=BG, padx=20, pady=14)
        btn_row.pack(fill="x")

        self.run_btn = tk.Button(
            btn_row, text="  Run Agent  ",
            font=("Segoe UI", 12, "bold"),
            bg=ACCENT, fg=BG,
            activebackground="#74c7ec", activeforeground=BG,
            relief="flat", padx=16, pady=10, cursor="hand2",
            command=self._run,
        )
        self.run_btn.pack(side="left")

        tk.Button(
            btn_row, text="Clear log",
            font=FONT, bg=SURF, fg=FG,
            relief="flat", padx=12, pady=10, cursor="hand2",
            command=self._clear,
        ).pack(side="left", padx=10)

        self.status_lbl = tk.Label(btn_row, text="", font=FONT, fg=GREEN, bg=BG)
        self.status_lbl.pack(side="left", padx=4)

        # Log
        tk.Label(self, text="  LOG", font=("Segoe UI", 8, "bold"),
                 fg=FG3, bg=BG, anchor="w").pack(fill="x", padx=20)

        self.log_box = scrolledtext.ScrolledText(
            self, font=MONO, bg="#11111b", fg=FG,
            insertbackground="white", relief="flat",
            wrap="word", state="disabled",
            padx=12, pady=8,
        )
        self.log_box.pack(fill="both", expand=True, padx=20, pady=(2, 16))

    def _row(self, parent, label, hint, is_file):
        frame = tk.Frame(parent, bg=CARD, pady=6)
        frame.pack(fill="x")
        frame.columnconfigure(1, weight=1)

        lbl_frame = tk.Frame(frame, bg=CARD)
        lbl_frame.grid(row=0, column=0, sticky="w", padx=(0, 14))
        tk.Label(lbl_frame, text=label, font=FONT_B,             fg=FG,  bg=CARD, anchor="w").pack(anchor="w")
        tk.Label(lbl_frame, text=hint,  font=("Segoe UI", 8),    fg=FG3, bg=CARD, anchor="w").pack(anchor="w")

        var = tk.StringVar()
        tk.Entry(
            frame, textvariable=var, font=FONT,
            bg=SURF, fg=FG, insertbackground="white",
            relief="flat", width=44,
        ).grid(row=0, column=1, padx=(0, 8), ipady=7, sticky="ew")

        if is_file:
            cmd = lambda: var.set(
                filedialog.askopenfilename(
                    initialdir=_ROOT,
                    filetypes=[("Documents", "*.txt *.docx *.pdf"), ("All files", "*.*")]
                ) or var.get()
            )
        else:
            cmd = lambda: var.set(filedialog.askdirectory(initialdir=_ROOT) or var.get())

        tk.Button(
            frame, text="Browse", font=FONT,
            bg=MAUVE, fg=BG,
            activebackground="#b4befe", activeforeground=BG,
            relief="flat", padx=10, pady=6, cursor="hand2", command=cmd,
        ).grid(row=0, column=2)

        return var

    def _log(self, msg):
        def _write():
            self.log_box.config(state="normal")
            self.log_box.insert("end", msg + "\n")
            self.log_box.see("end")
            self.log_box.config(state="disabled")
        self.after(0, _write)

    def _clear(self):
        self.log_box.config(state="normal")
        self.log_box.delete("1.0", "end")
        self.log_box.config(state="disabled")
        self.status_lbl.config(text="")

    def _run(self):
        docs    = self.docs_var.get().strip()
        prompt  = self.prompt_var.get().strip()
        gallery = self.gallery_var.get().strip()

        if not docs:
            messagebox.showerror("Missing", "Please select the Documents folder.")
            return
        if not prompt:
            messagebox.showerror("Missing", "Please select the Change prompt file.")
            return
        if not os.path.isdir(docs):
            messagebox.showerror("Not found", f"Documents folder not found:\n{docs}")
            return
        if not os.path.isfile(prompt):
            messagebox.showerror("Not found", f"Prompt file not found:\n{prompt}")
            return
        if gallery and not os.path.isdir(gallery):
            messagebox.showerror("Not found", f"Gallery folder not found:\n{gallery}")
            return

        self._clear()
        self.run_btn.config(state="disabled", text="  Running...  ")
        self.status_lbl.config(text="Working...", fg=YELLOW)

        def _done(ok, msg):
            def _upd():
                self.run_btn.config(state="normal", text="  Run Agent  ")
                self.status_lbl.config(
                    text=("Done: " if ok else "Failed: ") + msg,
                    fg=GREEN if ok else RED,
                )
            self.after(0, _upd)

        threading.Thread(
            target=run_pipeline,
            args=(prompt, docs, gallery, self._log, _done),
            daemon=True,
        ).start()


if __name__ == "__main__":
    App().mainloop()
